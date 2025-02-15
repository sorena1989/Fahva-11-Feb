import streamlit as st
import pandas as pd
import openai
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import os
import re
import tempfile
import zipfile
from io import BytesIO
from docx.shared import Pt
import requests
from bs4 import BeautifulSoup
from googlesearch import search
from urllib.parse import urlparse

# ---------------------------
# تنظیم API Key به صورت ثابت
# ---------------------------
OPENAI_API_KEY = "API Key" 
if not OPENAI_API_KEY:
    st.error("کلید API تنظیم نشده است!")
    st.stop()

openai.api_key = OPENAI_API_KEY 

# ---------------------------
# تابع ورود (Login)
# ---------------------------
def login():
    # اجرای ورود در ابتدای برنامه
    if "logged_in" not in st.session_state:
        st.session_state["logged_in"] = False
    if not st.session_state["logged_in"]:
        st.sidebar.subheader("ورود به داشبورد")
        username = st.sidebar.text_input("نام کاربری")
        password = st.sidebar.text_input("رمز عبور", type="password")
        if st.sidebar.button("ورود"):
            if username == "user" and password == "pass":
                st.session_state["logged_in"] = True
                st.sidebar.success("ورود موفقیت‌آمیز!")
                st.experimental_rerun()  # رفرش بلافاصله پس از ورود موفق
            else:
                st.sidebar.error("نام کاربری یا رمز عبور اشتباه است.")
    return st.session_state["logged_in"]

# ---------------------------
# توابع کمکی استخراج اطلاعات از اکسل
# ---------------------------

def extract_keywords_from_row(row):
    separator_pattern = r"[،;|؛]+"
    keywords_main = re.split(separator_pattern, str(row.get("کلمه کلیدی اصلی", "") or ""))
    keywords_main = [kw.strip() for kw in keywords_main if kw.strip()]
    keywords_sub = re.split(separator_pattern, str(row.get("کلمات کلیدی فرعی", "") or ""))
    keywords_sub = [kw.strip() for kw in keywords_sub if kw.strip()]
    return keywords_main, keywords_sub

def extract_prompt_data(row):
    original_word_count = row.get("تعداد کلمه", 2000)
    double_word_count = original_word_count * 1  # در صورت نیاز می‌توانید این ضریب را تغییر دهید
    return {
        "موضوع": row.get("موضوع", "موضوع مشخص نشده"),
        "نوع مقاله": row.get("نوع مقاله", "نوع مقاله مشخص نشده"),
        "تعداد کلمه": double_word_count,
        "کلمات کلیدی": row.get("کلمات کلیدی", None),
        "لینک‌ها": [row.get("لینک1"), row.get("لینک2")],
        "انکرتکست": [row.get("انکرتکست1"), row.get("انکرتکست2")],
        "عنوان اصلی": row.get("عنوان اصلی(H1)", "عنوان خلاقانه مقاله"),
        "عناوین H2": row.get("عناوین H2", None),
        "عناوین H3": row.get("عناوین H3", None)
    }

def remove_urls(text):
    return re.sub(r'https?://\S+', '', text)

# ---------------------------
# تابع تولید پرامپت پویا (شامل تمام دستورالعمل‌های اصلی)
# ---------------------------
def generate_dynamic_prompt(data, keywords_main, keywords_sub, crawled_info):
    """تولید پرامپت پویا بر اساس اطلاعات اکسل و اطلاعات کراول شده"""
    if data["نوع مقاله"] == "سفرنامه":
        prompt = f"لطفاً یک سفرنامه فارسی برای موضوع زیر بنویس:\n"
        prompt += f"- موضوع: {data['موضوع']}\n"
        prompt += f"- تعداد کلمات: حدود {data['تعداد کلمه']}\n"
        if keywords_main:
            prompt += f"- کلمات کلیدی اصلی: {', '.join(keywords_main)}\n"
        if keywords_sub:
            prompt += f"- کلمات کلیدی فرعی: {', '.join(keywords_sub)}\n"
        if any(data["لینک‌ها"]):
            prompt += "- لینک‌ها:\n"
            for link in data["لینک‌ها"]:
                if link:
                    prompt += f"  - {link}\n"
        if crawled_info:
            prompt += "\nاطلاعات استخراج شده از اینترنت:\n"
            prompt += f"{crawled_info}\n"
        prompt += "\nنکات راهنما:\n"
        prompt += "- متن سفرنامه باید از زبان اول شخص نوشته شود و لحن آن صمیمی و جذاب باشد.\n"
        prompt += "- تجربیات شخصی، تصمیم‌گیری‌ها و اتفاقات روزانه سفر را بیان کن.\n"
        prompt += "- در متن سفرنامه از هر کلمه کلیدی اصلی حداقل 3 بار و از هر کلمه کلیدی فرعی حداقل 2 بار به صورت طبیعی استفاده کن.\n"
        prompt += "- برای هر عنوان H1 حداقل دو پاراگراف طولانی و جذاب بنویس.\n"
        prompt += "- برای هر عنوان H2 و H3 حداقل یک پاراگراف خیلی طولانی (حداقل 5 سطر) آماده کن.\n"
        prompt += "- عناوین H2 و H3 باید منطقی و مرتبط بوده و بین آن‌ها یک متن مرتبط وجود داشته باشد.\n"
        prompt += "- در توضیحات هر عنوان از جزئیات مرتبط مانند اسامی افراد، مکان‌ها، تاریخچه، سبک معماری و غیره استفاده کن.\n"
        prompt += "- لینک‌ها باید کلیک‌پذیر بوده و به صورت طبیعی در متن قرار گیرند.\n"
        prompt += "- از عبارات بولد استفاده نکن و از شماره‌گذاری عناوین پرهیز کن.\n"
        prompt += "- اطلاعات دقیق و مستند ارائه کن.\n"
        prompt += "- به هیچ عنوان عبارات با ساختار (هتل'نام شهر') را به عنوان هتل واقعی در نظر نگیر.\n"
        prompt += "- بخش نتیجه‌گیری باید با عنوانی خلاقانه و متناسب با موضوع بوده و از عبارات کلیشه‌ای اجتناب شود.\n"
        prompt += "- نوع نگارش: \"Human-like, engaging, detailed, SEO-friendly\"\n"
        prompt += "- عناوین باید با تعداد `#`های مناسب شروع شده و بعد از آنها یک فاصله وجود داشته باشد.\n"
    else:
        prompt = f"لطفاً یک مقاله فارسی برای موضوع زیر بنویس:\n"
        prompt += f"- موضوع: {data['موضوع']}\n"
        prompt += f"- نوع مقاله: {data['نوع مقاله']}\n"
        prompt += f"- تعداد کلمات: حدود {data['تعداد کلمه']}\n"
        if keywords_main:
            prompt += f"- کلمات کلیدی اصلی: {', '.join(keywords_main)}\n"
        if keywords_sub:
            prompt += f"- کلمات کلیدی فرعی: {', '.join(keywords_sub)}\n"
        if any(data["لینک‌ها"]):
            prompt += "- لینک‌ها:\n"
            for link in data["لینک‌ها"]:
                if link:
                    prompt += f"  - {link}\n"
        if crawled_info:
            prompt += "\nاطلاعات استخراج شده از اینترنت:\n"
            prompt += f"{crawled_info}\n"
        prompt += "\nنکات راهنما:\n"
        prompt += "- لحن مقاله باید متناسب با نوع مقاله باشد.\n"
        prompt += "- از هر کلمه کلیدی اصلی حداقل 3 بار و از هر کلمه کلیدی فرعی حداقل 2 بار به صورت طبیعی استفاده کن.\n"
        prompt += "- برای هر عنوان H1 حداقل دو پاراگراف طولانی و جذاب بنویس.\n"
        prompt += "- برای هر عنوان H2 و H3 حداقل یک پاراگراف خیلی طولانی (حداقل 5 سطر) آماده کن.\n"
        prompt += "- هر H2 باید حداقل بین 2 تا 4 H3 داشته باشد.\n"
        prompt += "- لینک‌ها باید کلیک‌پذیر بوده و به صورت طبیعی در متن قرار گیرند.\n"
        prompt += "- اطلاعات دقیق و مستند ارائه کن و از گمانه‌زنی پرهیز کن.\n"
        prompt += "- از عبارات بولد استفاده نکن و از شماره‌گذاری عناوین خودداری کن.\n"
        prompt += "- به هیچ عنوان عبارات با ساختار (هتل'نام شهر') را به عنوان هتل واقعی در نظر نگیر.\n"
        prompt += "- بخش نتیجه‌گیری باید با عنوانی خلاقانه و متناسب با موضوع بوده و از عبارات کلیشه‌ای اجتناب کند.\n"
        prompt += "- تمام اصول SEO رعایت شود.\n"
        prompt += "- نوع نگارش: \"Human-like, engaging, detailed, SEO-friendly\"\n"
        prompt += "- عناوین باید با تعداد `#`های مناسب شروع شده و بعد از آنها یک فاصله وجود داشته باشد.\n"
    return prompt

# ---------------------------
# تابع تولید مقاله توسط OpenAI
# ---------------------------
def generate_article(prompt, model_name="gpt-4o-mini"):
    openai.api_key = OPENAI_API_KEY
    response = openai.ChatCompletion.create(
        model=model_name,
        messages=[
            {"role": "system", "content": (
                "شما یک نویسنده حرفه‌ای و دقیق فارسی هستید که مقالات جامع، خلاقانه و با جزئیات دقیق تولید می‌کنید. "
                "اطلاعات ارائه شده باید دقیق و مستند بوده و از الگوی نگارش SEO-friendly، زبان انسانی و جذاب استفاده شود."
            )},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
        max_tokens=4000
    )
    return response.choices[0].message.content



def calculate_word_count(text):
    words = re.findall(r'\b\w+\b', text.replace('\u200c', ''))
    return len(words)

def reduce_word_count_naturally(article_content, desired_word_count):
    revise_prompt = (
        f"مقاله زیر باید به تعداد کلمات حدود {desired_word_count} کاهش یابد.\n\n"
        f"--- متن مقاله فعلی ---\n"
        f"{article_content}\n"
        f"-----------------------\n"
        f"لطفاً متن مقاله را به صورت طبیعی و بدون از دست دادن اطلاعات مهم کاهش دهید تا تعداد کلمات آن به حدود {desired_word_count} برسد."
    )
    try:
        reduced_content = generate_article(revise_prompt)
        return reduced_content
    except Exception as e:
        st.error(f"خطا در کاهش تعداد کلمات مقاله: {e}")
        return article_content

# ---------------------------
# افزودن لینک به پاراگراف در فایل Word
# ---------------------------
def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    new_text = OxmlElement('w:t')
    new_text.text = text
    new_run.append(new_text)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)

# ---------------------------
# ذخیره مقاله در فایل Word با پشتیبانی از چندین لینک در یک خط
# ---------------------------
def save_to_word(content, filename, links):
    doc = Document()
    # تنظیم استایل اصلی برای راست به چپ
    doc.styles['Normal'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.styles['Normal']._element.get_or_add_pPr().set(qn('w:bidi'), '1')
    
    if links:
        pattern = "(" + "|".join(re.escape(key) for key in links.keys()) + ")"
    else:
        pattern = None
    
    lines = content.split('\n')
    for line in lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph._element.get_or_add_pPr().set(qn('w:bidi'), '1')
        if pattern and re.search(pattern, line):
            parts = re.split(pattern, line)
            for part in parts:
                if part in links:
                    add_hyperlink(paragraph, part, links[part])
                else:
                    paragraph.add_run(part)
        else:
            paragraph.add_run(line)
        p_format = paragraph.paragraph_format
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        p_format.line_spacing = 1.5
    doc.save(filename)

def check_keywords_in_text(keywords, text):
    missing_keywords = [kw.strip() for kw in keywords if kw.strip() not in text]
    return missing_keywords

# ---------------------------
# اعتبارسنجی ساختار مقاله
# ---------------------------
def validate_article_structure(content, data):
    """بررسی ساختار مقاله جهت اطمینان از وجود متن بین عناوین"""
    lines = content.split('\n')
    previous_heading = None
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('## '):
            previous_heading = 'H2'
        elif stripped.startswith('### '):
            if previous_heading not in ['H2', 'H3', 'Text']:
                return False
            previous_heading = 'H3'
        elif stripped and not stripped.startswith('#'):
            previous_heading = 'Text'
    return True

# ---------------------------
# توابع مربوط به کراولینگ اطلاعات
# ---------------------------
def crawl_google_links(topic, num_links=5):
    links = []
    try:
        query = topic
        count = 0
        for url in search(query):
            count += 1
            if "aclk" in url or "sponsored" in url:
                continue
            if url not in links:
                links.append(url)
            if len(links) >= num_links or count >= 20:
                break
    except Exception as e:
        st.error(f"خطا در جستجوی گوگل: {e}")
    return links


def fetch_page_content(url):
    try:
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, "html.parser")
            article = soup.find("article")
            if article:
                text = article.get_text(separator="\n")
            else:
                text = soup.get_text(separator="\n")
            return text
        else:
            return ""
    except Exception as e:
        st.error(f"خطا در دریافت محتوای {url}: {e}")
        return ""

def gather_crawled_information(topic):
    links = crawl_google_links(topic, num_links=5)
    if not links:
        st.warning("هیچ نتیجه‌ای یافت نشد.")
        return ""
    info_list = []
    for link in links:
        content = fetch_page_content(link)
        if content:
            info_list.append(f"منبع: {link}\n{content}")
    crawled_info = "\n\n".join(info_list)
    crawled_info_clean = remove_urls(crawled_info)
    display_text = crawled_info_clean[:500] + "..." if len(crawled_info_clean) > 500 else crawled_info_clean
    with st.expander("نمایش بخشی از اطلاعات کراول شده"):
        st.text(display_text)
    return crawled_info_clean

# ---------------------------
# اپلیکیشن Streamlit
# ---------------------------
def main():
    if not login():
        st.stop()
        
    st.set_page_config(page_title="Fahva Article Generator Beta V1.1", layout="wide")
    st.title("Fahva Article Generator alltour V1.1")
    
    st.sidebar.header("تنظیمات")
    model_name = st.sidebar.selectbox("انتخاب مدل OpenAI", ["gpt-4o-mini", "gpt-3.5-turbo"])
    uploaded_file = st.sidebar.file_uploader("آپلود فایل اکسل", type=["xlsx", "xls"])
    
    if st.sidebar.button("شروع تولید مقالات"):
        if not uploaded_file:
            st.error("لطفاً فایل اکسل را آپلود کنید.")
            return
        try:
            df = pd.read_excel(uploaded_file, engine='openpyxl')
        except Exception as e:
            st.error(f"خطا در خواندن فایل اکسل: {e}")
            return
        required_columns = ["موضوع", "نوع مقاله", "تعداد کلمه", "کلمه کلیدی اصلی", "کلمات کلیدی فرعی", "لینک1", "انکرتکست1"]
        for col in required_columns:
            if col not in df.columns:
                st.error(f"ستون اجباری '{col}' در فایل اکسل وجود ندارد.")
                return
        with tempfile.TemporaryDirectory() as tmpdirname:
            st.success(f"مسیر موقت برای ذخیره فایل‌ها: {tmpdirname}")
            generated_files = []
            for idx, row in df.iterrows():
                data = extract_prompt_data(row)
                keywords_main, keywords_sub = extract_keywords_from_row(row)
                st.info(f"در حال جستجو و کراول اطلاعات مرتبط با موضوع '{data['موضوع']}' برای ردیف {idx + 1}...")
                crawled_info = gather_crawled_information(data["موضوع"])
                attempt = 0
                max_attempts = 3
                successful = False
                while attempt < max_attempts and not successful:
                    attempt += 1
                    st.write(f"تولید مقاله برای ردیف {idx + 1} - تلاش {attempt}")
                    prompt = generate_dynamic_prompt(data, keywords_main, keywords_sub, crawled_info)
                    try:
                        article_content = generate_article(prompt, model_name)
                        missing_main = check_keywords_in_text(keywords_main, article_content) if keywords_main else []
                        missing_sub = check_keywords_in_text(keywords_sub, article_content) if keywords_sub else []
                        desired_word_count = int(data['تعداد کلمه'])
                        lower_limit = int(desired_word_count * 0.9)
                        upper_limit = int(desired_word_count * 1.1)
                        word_count = calculate_word_count(article_content)
                        
                        if word_count > upper_limit:
                            st.info(f"تعداد کلمات ({word_count}) بیش از حد مجاز ({upper_limit}) است. در حال کاهش تعداد کلمات...")
                            article_content = reduce_word_count_naturally(article_content, desired_word_count)
                            word_count = calculate_word_count(article_content)
                            st.info(f"تعداد کلمات پس از کاهش: {word_count}")
                        if word_count < lower_limit:
                            missing_keywords = []
                            if missing_main:
                                missing_keywords.append(f"کلمات کلیدی اصلی: {', '.join(missing_main)}")
                            if missing_sub:
                                missing_keywords.append(f"کلمات کلیدی فرعی: {', '.join(missing_sub)}")
                            revise_prompt = (
                                f"مقاله زیر باید بازنویسی شود تا:\n"
                                f"- تعداد کلمات به حدود {desired_word_count} برسد.\n"
                                f"- کلمات کلیدی زیر که استفاده نشده‌اند، به متن اضافه شوند:\n"
                                f"{'; '.join(missing_keywords)}\n\n"
                                f"--- متن مقاله فعلی ---\n"
                                f"{article_content}\n"
                                f"-----------------------\n"
                                f"لطفاً متن مقاله را گسترش دهید و کلمات کلیدی بالا را به طور طبیعی و مناسب در متن بگنجانید."
                            )
                            article_content = generate_article(revise_prompt, model_name)
                            word_count = calculate_word_count(article_content)
                            st.info(f"تعداد کلمات پس از بازنویسی: {word_count}")
                        
                        # اعتبارسنجی ساختار مقاله قبل از ذخیره
                        if not validate_article_structure(article_content, data):
                            st.warning(f"ساختار مقاله ردیف {idx + 1} ممکن است ناقص باشد.")
                        
                        if missing_main or missing_sub or not (lower_limit <= word_count <= upper_limit):
                            st.warning(f"مقاله ردیف {idx + 1} نیاز به بازنویسی دارد.")
                            if missing_main:
                                st.warning(f"کلمات کلیدی اصلی استفاده نشده: {', '.join(missing_main)}")
                            if missing_sub:
                                st.warning(f"کلمات کلیدی فرعی استفاده نشده: {', '.join(missing_sub)}")
                        else:
                            successful = True
                            st.success(f"مقاله ردیف {idx + 1} با موفقیت تولید شد.")
                            st.info(f"تعداد کلمات نهایی مقاله: {word_count}")
                            safe_title = re.sub(r'[\\/*?:"<>|]', "", data['عنوان اصلی']).strip().replace(' ', '_')
                            # استفاده از لینک‌های موجود در فایل اکسل (در صورت وجود)
                            file_links = {f"لینک{i+1}": link for i, link in enumerate(data["لینک‌ها"]) if link}
                            output_filename = os.path.join(tmpdirname, f"article_{safe_title}_{idx + 1}.docx")
                            save_to_word(article_content, output_filename, file_links)
                            generated_files.append((f"مقاله {idx + 1}: {data['عنوان اصلی']}", output_filename))
                            st.write(f"مقاله ذخیره شد: {output_filename}")
                    except Exception as e:
                        st.error(f"خطا در تولید مقاله برای ردیف {idx + 1}: {e}")
                if not successful:
                    st.error(f"مقاله ردیف {idx + 1} پس از {max_attempts} تلاش تولید نشد.")
            if generated_files:
                zip_buffer = BytesIO()
                with zipfile.ZipFile(zip_buffer, "w") as zipf:
                    for file_name, file_path in generated_files:
                        zipf.write(file_path, os.path.basename(file_path))
                zip_buffer.seek(0)
                st.download_button(
                    label="دانلود همه مقالات به صورت ZIP",
                    data=zip_buffer,
                    file_name="generated_articles.zip",
                    mime="application/zip"
                )
            else:
                st.warning("هیچ مقاله‌ای تولید نشد.")

if __name__ == "__main__":
    main()
